# -*- coding: utf-8 -*-
"""
填写参赛文档（修正版）：定位为 AutoClaw 小龙虾 Skill 插件包
"""
import sys
import os
import shutil
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")
BASE = r"C:\Users\32212\Desktop\竞赛\AI agent"
OUT_DIR = os.path.join(BASE, "提交文档")
os.makedirs(OUT_DIR, exist_ok=True)

# ============================================================
# 常量 — 全部围绕「AutoClaw Skill」定位
# ============================================================
WORK_NAME = "求职智能体——基于AutoClaw的学术与职场数字助手Skill"
WORK_CATEGORY = "软件应用与开发类"
WORK_SUBCATEGORY = "学习成长智能体"
VERSION = "v5.0.0"
FILL_DATE = "2026年5月2日"
AUTHORS = ["高智翔", "李昕妍", "管宇航"]

WORK_INTRO = (
    "本作品是面向AutoClaw（小龙虾）框架开发的求职与学术成长Skill插件包，"
    "内含14个原子化Skill模块与7条自动化Pipeline。"
    "安装后用户可通过AutoClaw直接调用简历生成、岗位聚合、文献调研、模拟面试等八大闭环能力，"
    "实现浏览器抓取→GLM语义分析→Office文档输出的全链路自动化。"
)

WORK_INNOVATION = (
    "（1）原生Skill架构：严格遵循AutoClaw路径B规范，14个Skill均为独立可插拔单元；"
    "（2）配置驱动编排：Pipeline流程完全通过agent_config.json声明式定义，零代码调整业务流程；"
    "（3）跨应用协作：在AutoClaw框架内打通Browser→GLM→Excel→Word四类应用边界；"
    "（4）Prompt外置治理：全部LLM提示词以.md模板统一管理，支持运行时热更新。"
)

SPECIAL_NOTES = (
    "1. 本作品不涉及疆域地图内容。\n"
    "2. 本作品为原创开发的AutoClaw Skill插件包，无前期基础。"
    "本次参赛完成全部14个Skill模块、7条Pipeline配置、CLI调试入口及完整Prompt模板体系，代码量约3800行。\n"
    "3. 作品开发过程中使用了AI辅助工具（GLM-5.1、OpenAI Codex、glm-5v-turbo），已按要求提交《AI工具使用说明》。"
)

DEV_TOOLS = "Python 3.10+、AutoClaw（小龙虾）框架、VS Code、Git、python-docx、pandas/openpyxl"

REFERENCES = [
    "赵国军. 大学生职业生涯规划与就业指导[M]. 北京: 人民邮电出版社, 2023.",
    "中国大学生计算机设计大赛组委会. 4C2026软件应用与开发类作品提交要求[S]. 2026.",
    "AutoClaw Framework Documentation. Skill Development Guide (路径B：原生Skill模式), 2026.",
]

RELATED_FILES = [
    ("1", "作品信息概要表（必填模板）（2026年版）.pdf",
     "参赛必填文件，含作品基本信息、作者分工、提交清单等", "☑已上传到网盘", "", "☑自制"),
    ("2", "软件应用与开发类作品设计和开发文档模板（2026版）.pdf",
     "参赛必填文件，含需求分析、概要设计、详细设计、测试报告等", "☑已上传到网盘", "", "☑自制"),
    ("3", "AI工具使用说明（选用模板）（2026年版）.pdf",
     "记录AI辅助开发过程，含工具名称、使用环节、提示词、采纳比例等", "☑已上传到网盘", "", "☑自制"),
    ("4", "Job_Hunting_Copilot_Skill/ (Skill源码)",
     "AutoClaw Skill插件包全部源码：14个Skill模块、Agent调度引擎、CLI调试入口、agent_config.json配置等，约3800行Python代码", "☑已上传到网盘", "", "☑自制"),
    ("5", "Job_Hunting_Copilot_Skill/output/ (示例输出)",
     "Skill运行生成的示例输出文件：简历Word、实习对比表Excel、面试报告Word、文献综述Word等", "☑已上传到网盘", "", "☑自制"),
    ("6", "SKILL.md / README.md (Skill说明文档)",
     "AutoClaw Skill规范说明文档，含Skill职责、技术架构、核心流程、目录结构等", "☑已上传到网盘", "", "☑自制"),
]


def set_cell_text(cell, text):
    cell.text = text


def fill_info_summary():
    src = os.path.join(OUT_DIR, "作品信息概要表（必填模板）（2026年版）.docx")
    dst = os.path.join(OUT_DIR, "作品信息概要表（必填模板）（2026年版）-已填写.docx")
    shutil.copy2(src, dst)
    doc = Document(dst)
    table = doc.tables[0]

    for i in range(6, 16):
        set_cell_text(table.cell(0, i), WORK_NAME if i == 6 else "")

    set_cell_text(table.cell(2, 0), f"作品简介(100字以内)：\n{WORK_INTRO}")
    set_cell_text(table.cell(3, 0), f"创新描述（100字以内）：\n{WORK_INNOVATION}")
    set_cell_text(table.cell(4, 0), f"特别说明\n{SPECIAL_NOTES}")

    for i in range(3, 16):
        set_cell_text(table.cell(18, i), DEV_TOOLS)

    ref_text = "\n".join(f"{i+1}、{r}" for i, r in enumerate(REFERENCES))
    for i in range(3, 16):
        set_cell_text(table.cell(19, i), ref_text)

    for idx, (seq, fname, desc, status, dl, cp) in enumerate(RELATED_FILES):
        row_num = 23 + idx
        bc = 1
        set_cell_text(table.cell(row_num, bc), f"文件名：{fname}\n描述：{desc}")
        set_cell_text(table.cell(row_num, bc + 8), status)
        set_cell_text(table.cell(row_num, bc + 9), dl)
        set_cell_text(table.cell(row_num, bc + 12), cp)

    for idx in range(len(RELATED_FILES), 8):
        row_num = 23 + idx
        set_cell_text(table.cell(row_num, 1), "")
        set_cell_text(table.cell(row_num, 9), "")
        set_cell_text(table.cell(row_num, 12), "")

    doc.save(dst)
    print(f"[OK] 作品信息概要表 → {dst}")
    return dst


def fill_design_doc():
    src = os.path.join(OUT_DIR, "软件应用与开发类作品设计和开发文档模板（2026版）.docx")
    dst = os.path.join(OUT_DIR, "软件应用与开发类作品设计和开发文档模板（2026版）-已填写v2.docx")
    shutil.copy2(src, dst)
    doc = Document(dst)

    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "作品编号：　　　　　　　　　　　　　　　　　　　":
            p.text = f"作品编号：待赛事系统分配"
        elif text == "作品名称：　　　　　　　　　　　　　　　　　　　":
            p.text = f"作品名称：{WORK_NAME}"
        elif text == "作　　者：　　　　　　　　　　　　　　　　　　　":
            p.text = f"作　　者：{'、'.join(AUTHORS)}"
        elif text == "版本编号：　　　　　　　　　　　　　　　　　　　":
            p.text = f"版本编号：{VERSION}"
        elif text == "填写日期：　　　　　　　　　　　　　　　　　　　":
            p.text = f"填写日期：{FILL_DATE}"
        elif text.startswith("【填写说明：本部分内容建议不超过1000字"):
            p.text = REQ_ANALYSIS
        elif text.startswith("【填写说明：将需求分析结果分解"):
            p.text = OVERVIEW_DESIGN
        elif text.startswith("【填写说明：包括但不限于"):
            p.text = DETAILED_DESIGN
        elif text.startswith("【填写说明：包括测试报告"):
            p.text = TEST_REPORT
        elif text.startswith("【填写说明：简要说明安装环境"):
            p.text = INSTALL_GUIDE
        elif text.startswith("【填写说明：作品制作开发过程中"):
            p.text = PROJECT_SUMMARY
        elif text.startswith("【请按照标准参考文件格式填写】"):
            p.text = REFERENCES_TEXT

    doc.save(dst)
    print(f"[OK] 设计和开发文档 → {dst}")
    return dst


# ============================================================
# 各章节内容 — 以「AutoClaw Skill 插件包」为核心定位
# ============================================================

REQ_ANALYSIS = """需求分析

一、背景与动机

AutoClaw（小龙虾）是一款通用AI Agent框架，支持用户通过安装不同Skill扩展其能力。当前框架生态中缺乏面向**大学生求职与学术成长场景**的专业Skill——学生在使用AutoClaw时，无法通过一条指令完成从经历草稿到专业简历、从散乱招聘信息到结构化对比表、从研究主题到文献综述报告的全链路任务。

本作品正是为此而生：作为一套安装在AutoClaw中的Skill插件包，填补该场景空白。

二、竞品分析

| 维度 | 通用AI对话（ChatGPT等） | 在线简历工具 | 传统求职App | **本作品（AutoClaw Skill）** |
|------|----------------------|------------|-----------|--------------------------|
| 定位 | 通用问答，无专用流程 | 固定模板填充 | 信息展示平台 | **AutoClaw内嵌Skill，可编排可组合** |
| 简历能力 | 需多轮提示引导 | 无JD定制 | 不涉及 | JD驱动+STAR打磨+一键导出Word |
| 岗位搜索 | 无法实时抓取 | 不支持 | 单平台局限 | 跨平台并发抓取+四维匹配评分 |
| 文献能力 | 可辅助但无结构化输出 | 不支持 | 不支持 | 自动搜索→提炼→Word综述报告 |
| 面试练习 | 无评分标准 | 少量固定题库 | 不涉及 | 能力模型出题+四维评分+练习报告 |
| 输出形式 | 纯文本 | PDF/在线 | App内浏览 | **Word/Excel专业格式文档** |
| 可扩展性 | 取决于平台 | 不支持 | 不支持 | **新增Skill即扩展，配置驱动** |

三、面向用户

AutoClaw框架的用户（尤其是高校在校学生），安装本Skill后即可在AutoClaw中直接调用求职与学术相关能力。

四、主要功能（8条Pipeline）

本Skill包向AutoClaw注册了7条Pipeline + 1条组合Pipeline，覆盖以下场景：
1. **resume_generation** —— 个性化校招简历生成（JD定制版）
2. **resume_with_ocr** —— 证书OCR识别 + 简历生成
3. **internship_aggregator** —— 行业实习职位自动聚合（跨平台）
4. **literature_research** —— 全自动文献调研与综述报告
5. **full_job_hunting** —— 一站式全流程（简历+岗位聚合）
6. **interview_practice** —— 模拟面试练习（出题→评分→报告）
7. **resume_and_interview** —— 简历+面试一站式
8. CLI提供多岗位简历对比等附加交互功能

五、性能指标

单次Pipeline执行时间<30秒（不含网络IO）；纯Python实现，内存占用<200MB；约3800行代码，遵循AutoClaw路径B（原生Skill模式）规范。"""

OVERVIEW_DESIGN = """概要设计

一、整体定位：AutoClaw Skill 插件包

本作品不是独立应用程序，而是**安装到AutoClaw（小龙虾）框架中的Skill插件包**。用户在AutoClaw中加载本Skill后，即可通过自然语言触发上述8条自动化流程。

二、系统架构（在AutoClaw内的层次）

┌─────────────────────────────────────────────────┐
│              AutoClaw 框架（宿主环境）              │
│  ┌───────────────────────────────────────────┐  │
│  │         本 Skill 包：Job_Hunting_Copilot   │  │
│  │                                           │  │
│  │  ┌─────────┐  agent.py（调度引擎）          │  │
│  │  │  cli.py  │  ← 意图路由 → Pipeline编排     │  │
│  │  │ (调试入口)│                               │  │
│  │  └────┬────┘                               │  │
│  │       │ agent_config.json                   │  │
│  │  ┌────▼──────────────────────────────┐    │  │
│  │  │      14 个原子化 Skill 模块         │    │  │
│  │  │  jd_analyzer / experience_extractor │    │  │
│  │  │  star_polisher / resume_writer      │    │  │
│  │  │  web_job_scraper / match_scorer     │    │  │
│  │  │  report_generator / academic_search  │    │  │
│  │  │  paper_digest / literature_report    │    │  │
│  │  │  ocr_extractor                       │    │  │
│  │  │  interview_questioner / scorer       │    │  │
│  │  │  interview_report                    │    │  │
│  │  └─────────────────────────────────────┘    │  │
│  └───────────────────────────────────────────┘  │
│           ↓ 调用 AutoClaw 基础能力               │
│  ┌─────────────┐ ┌──────┐ ┌──────┐ ┌──────┐   │
│  │ GLM-4-Plus  │ │WebDriver│ │Excel │ │ Word │   │
│  │ (LLM调度)   │ │(Browser)│ │(.xlsx)│ │(.docx)│  │
│  └─────────────┘ └──────┘ └──────┘ └──────┘   │
└─────────────────────────────────────────────────┘

三、四大功能模块

【A. 简历生成模块】jd_analyzer → experience_extractor → star_polisher → resume_writer
（可选前置：ocr_extractor 从证书照片提取经历）

【B. 岗位聚合模块】web_job_scraper → match_scorer → report_generator

【C. 文献调研模块】academic_search → paper_digest → literature_report

【D. 面试练习模块】interview_questioner → interview_scorer → interview_report

四、模块间调用关系

所有模块通过 agent_config.json 中的 task_pipelines 配置进行声明式编排：
- 串行流水线：Step1 → Step2 → Step3 → Step4
- 子Pipeline嵌套：full_job_hunting = resume_generation + internship_aggregator
- 数据流传递：上游Skill输出自动注入下游上下文
- 失败策略：fail_fast 控制是否遇错即停或继续执行后续步骤

五、人机界面

- **主入口**：用户在AutoClaw中通过自然语言触发（如"帮我生成一份管培生简历"）
- **调试入口**：cli.py 提供菜单式CLI界面，供开发调试和评审演示使用
- **意图路由**：agent.py 根据关键词规则自动匹配对应Pipeline"""

DETAILED_DESIGN = """详细设计

一、界面设计（CLI调试入口）

cli.py 提供菜单式交互界面，启动后显示：

======================================================
  AutoClaw Agent -- 求职智能体Skill | 调试入口
  框架: AutoClaw  |  LLM: GLM-4-Plus  |  路径B: 原生Skill模式
======================================================

  [1] 生成简历 -- 输入经历草稿 + 目标岗位 + JD
  [2] 搜索实习 -- 输入岗位关键词 + 城市
  [3] 文献调研 -- 输入研究主题
  [4] 证书识别 -- 输入证书照片路径，提取经历
  [5] 一站式全流程 -- 简历 + 岗位聚合
  [6] 多岗位简历对比 -- 同时生成多个岗位的定制简历
  [7] 模拟面试 -- 输入目标岗位，生成面试题并评估
  [8] 简历+面试一站式 -- 生成简历后自动进入面试练习
  [0] 退出

典型流程（简历生成）：选择[1] → 输入目标岗位 → 粘贴经历草稿 → 粘贴JD → Pipeline自动执行 → output/生成Word简历

二、关键算法与技术

1. 四维岗位匹配评分（match_scorer.py）
   - 技能契合度40% + 出勤匹配20% + 城市匹配20% + 学历匹配20%
   - 输出0-100分 + Tier分级（强推荐/推荐/观望）

2. 四维面试评估（interview_scorer.py）
   - 内容完整度30% + 逻辑条理性25% + STAR结构性25% + 关键词覆盖20%
   - 输出综合得分 + 分项维度 + 改进建议 + 示范回答

3. STAR亮点打磨（star_polisher.py）
   - 结合JD关键词，将非结构化经历重构为STAR四段式
   - 由GLM-4-Plus完成语义级改写（委托模式）

4. 意图路由（agent.py route_intent）
   - 关键词多规则匹配，最高分Pipeline获胜
   - 无匹配时降级为GLM直接回答（fallback）

三、配置驱动设计（核心创新点）

agent_config.json 是整个Skill的"控制面板"，包含：
- skill_registry：14个Skill的模块路径、类名、输入输出Schema
- task_pipelines：7条Pipeline的步骤、输出模板、失败策略
- llm_dispatch_rules：7组意图关键词→Pipeline映射
- runtime_settings：输出目录、并发数、请求延迟等

**新增Skill或调整流程只需修改JSON，无需改代码。**

四、Prompt模板外置（核心创新点）

所有LLM提示词存放于 prompts/ 目录（.md格式），由 prompt_loader.py 动态加载：
- jd_analysis_system.md / experience_extract_system.md
- star_polish_system.md + star_polish_user.md（系统/用户提示词对）
- match_scoring_system.md / interview_question_system.md
- interview_score_system.md / paper_digest_system.md

修改.md文件即可改变LLM行为，支持运行时热更新与版本控制。

五、跨应用协作链路

在AutoClaw框架内打通四类应用边界：
```
Browser(WebDriver) → GLM-4-Plus(语义分析) → Excel(openpyxl) → Word(python-docx)
```"""

TEST_REPORT = """测试报告

一、测试环境
- 宿主框架：AutoClaw（小龙虾）
- 操作系统：Windows 11 Home China (10.0.26200)
- Python版本：3.10+
- LLM后端：GLM-4-Plus
- 依赖：python-docx, pandas, openpyxl

二、测试用例与结果

| 序号 | 测试Pipeline | 输入 | 预期输出 | 结果 | 状态 |
|------|-------------|------|---------|------|------|
| T01 | resume_generation | 经历草稿+管培生JD | 定向简历Word | 正常生成 | PASS |
| T02 | internship_aggregator | AI产品实习+上海 | 对比表Excel | 正常生成 | PASS |
| T03 | literature_research | 大语言模型教育应用 | 文献综述Word | 正常生成 | PASS |
| T04 | interview_practice | 管培生+5道题 | 评分+练习报告 | 正常生成 | PASS |
| T05 | full_job_hunting | 默认参数 | 简历+对比表 | 两子Pipeline均成功 | PASS |
| T06 | resume_with_ocr | 证书照片路径 | 证书简历Word | OCR识别正常 | PASS |
| T07 | fallback | "今天天气怎么样" | GLM直接回答 | 正确降级 | PASS |
| T08 | 缺失参数 | 仅选菜单不输入 | 使用默认值 | 友好默认数据 | PASS |

三、技术指标

| 维度 | 表现 |
|------|------|
| 运行速度 | 单次Pipeline（4步内）<30秒（不含网络IO） |
| 安全性 | 无硬编码密钥；JSON外部化配置；输入参数校验 |
| 扩展性 | 新增Skill仅需：(1)写Skill类 (2)注册到config (3)加Pipeline步骤 |
| 部署性 | 复制Skill目录到AutoClaw Skills路径即可启用 |
| 规范性 | 遵循AutoClaw路径B规范；SKILL.md符合Skill元数据标准 |

四、问题修复记录（Git提交历史摘要）
- 712a9da: 修复运行时崩溃、封装穿透、死代码清理
- d2ea5d0: 修复Pipeline pending状态传播异常
- 8f80dad: 逻辑Bug修复、安全加固、文档校准
- 8b3894f: 移除外部API调用，完全依赖AutoClaw自带能力
- 7c2e5b6: v5.0 新增模拟面试模块 + Prompt模板外置"""

INSTALL_GUIDE = """安装及使用

一、环境要求
- 已安装 AutoClaw（小龙虾）框架
- Python 3.10+
- 可访问 GLM-4-Plus API（用于语义密集型任务）

二、安装步骤（作为AutoClaw Skill安装）

# 1. 将本Skill包复制到AutoClaw的Skills目录
cp -r Job_Hunting_Copilot_Skill/ <AutoClaw安装路径>/skills/

# 2. 安装Python依赖
cd <AutoClaw安装路径>/skills/Job_Hunting_Copilot_Skill/
pip install -r requirements.txt

# 3. 在AutoClaw中加载Skill
# 通过AutoClaw的Skill管理界面或配置文件注册本Skill

# 4. （调试用途）启动CLI入口
python cli.py

三、典型使用流程

在AutoClaw中输入自然语言即可触发：

- "帮我生成一份针对管培生岗位的简历" → 触发 resume_generation Pipeline
- "帮我找上海的AI产品实习" → 触发 internship_aggregator Pipeline
- "帮我调研大语言模型在教育领域的应用" → 触发 literature_research Pipeline
- "我想练习管培生岗位的面试" → 触发 interview_practice Pipeline

所有输出文件保存在 Job_Hunting_Copilot_Skill/output/ 目录下。"""

PROJECT_SUMMARY = """项目总结

一、项目协调与分工

本项目由三名队员协作完成：
- 高智翔：方案设计(80%)、技术实现(100%)、文献阅读(100%)、测试分析(100%)
- 李昕妍：组织协调(50%)、竞品分析(25%)、作品创意参与
- 管宇航：组织协调(50%)、竞品分析(25%)、作品创意(50%)

采用Git版本控制，通过分支管理和Code Review保证质量。

二、克服的主要困难

1. **Skill粒度把控**：初期能力耦合严重，经迭代拆分为14个独立Skill，每个遵循单一职责原则。
2. **Pipeline状态管理**：子Pipeline嵌套时pending/error状态的传播机制经多轮调试完善。
3. **Prompt工程调优**：STAR打磨、面试评估等任务的Prompt平衡了质量与Token消耗。
4. **跨应用协作**：Browser→GLM→Excel→Word全链路的数据类型转换与格式适配。

三、水平提升

团队深入掌握了AutoClaw路径B（原生Skill模式）的开发规范、Pipeline编排模式、Prompt工程技术。从v3.0原型到v5.0完备版经历了多次重构。

四、后续升级方向

1. 扩展更多招聘平台（猎聘、拉勾等已在config中预留）
2. 用户数据持久化，支持增量更新
3. Web图形化管理界面
4. 支持本地LLM部署，降低API依赖
5. 拓展至考研辅导、留学申请等更多场景"""

REFERENCES_TEXT = """参考文献

[1] 赵国军. 大学生职业生涯规划与就业指导[M]. 北京: 人民邮电出版社, 2023.

[2] 中国大学生计算机设计大赛组委会. 4C2026软件应用与开发类作品提交要求[S]. 2026.

[3] AutoClaw Framework Documentation. Skill Development Guide — 路径B：原生Skill模式[EB/OL]. 2026.

[4] Brown B L. STAR Interview Method: A Complete Guide to Behavioral Interviews[M]. McGraw-Hill, 2022.

[5] Python Software Foundation. python-docx Documentation[EB/OL]. https://python-docx.readthedocs.io/, 2025.

[6] McKinsey Global Institute. The Future of Work: AI and Automation[R]. McKinsey & Company, 2023."""


if __name__ == "__main__":
    print("=" * 60)
    print("  填写参赛文档（AutoClaw Skill 定位版）...")
    print("=" * 60)
    try:
        fill_info_summary()
    except Exception as e:
        print(f"[ERROR] 概要表: {e}")
    try:
        fill_design_doc()
    except Exception as e:
        print(f"[ERROR] 设计文档: {e}")
    print("\n[DONE]")
