# -*- coding: utf-8 -*-
"""生成「开源代码与组件使用情况说明」Word 文档（简短版）"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "开源代码与组件使用情况说明.docx")


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    HEADER_BG = "1F4E79"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # ── 标题 ──
    title = doc.add_heading("开源代码与组件使用情况说明", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()

    # ── 项目信息 ──
    h1 = doc.add_heading("一、项目信息", level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    add_table(doc,
              ["项目", "内容"],
              [
                  ["项目名称", "求职智能体（Job Hunting Copilot Skill）"],
                  ["框架", "AutoClaw"],
                  ["开发语言", "Python 3"],
                  ["LLM 后端", "智谱 AI GLM-4-Plus"],
              ],
              col_widths=[4, 12])

    # ── 开源组件 ──
    h2 = doc.add_heading("二、开源组件使用情况", level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    add_table(doc,
              ["组件名称", "版本", "许可证", "使用说明"],
              [
                  ["pandas", ">=1.5.0", "BSD-3-Clause",
                   "报告生成模块中用于 DataFrame 数据整理"],
                  ["openpyxl", ">=3.0.10", "MIT",
                   "Excel 文件格式化（条件格式、冻结行、自适应列宽）"],
                  ["python-docx", ">=0.8.11", "MIT",
                   "Word 简历文件 / 面试报告生成（页面布局、样式渲染）"],
              ],
              col_widths=[3, 2.5, 3, 7.5])

    # ── 合规声明 ──
    h3 = doc.add_heading("三、合规声明", level=1)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    run = p.add_run(
        "本项目所使用的第三方开源组件均采用 MIT、BSD-3-Clause 或 Apache 2.0 "
        "许可证，均为宽松型开源协议，允许商业使用、修改和分发。"
        "项目已保留所有原始版权声明和许可声明，符合各组件的许可证要求。"
    )
    run.font.size = Pt(10.5)

    # ── 保存 ──
    doc.save(OUTPUT_PATH)
    print(f"[OK] 文档已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
