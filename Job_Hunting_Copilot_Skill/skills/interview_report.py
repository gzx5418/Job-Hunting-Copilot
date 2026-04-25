"""
面试报告 Skill (Interview Report Skill)
AutoClaw Skill ID: interview_report

职责：
  汇总面试练习的问答与评分数据，
  生成 Word 面试练习报告（含评分雷达图数据、改进建议）。

触发场景：
  - 面试练习 Pipeline 中，评分完成后自动触发
  - 或当用户说"生成面试练习报告"

跨应用能力：
  GLM (评分数据) → Microsoft Word (.docx)
"""

import os
import re
import logging
from typing import List, Dict, Any
from skills import AutoClawSkill


class InterviewReportSkill(AutoClawSkill):
    """
    面试报告 Skill
    将面试练习数据生成 Word 报告
    """

    SKILL_NAME = "interview_report"
    SKILL_DESCRIPTION = "汇总面试问答与评分，生成 Word 练习报告（含维度分析、改进建议）。"

    def run(self, interview_results: List[Dict] = None,
            target_role: str = "未知岗位",
            output_path: str = "", **kwargs) -> Dict[str, Any]:
        """
        生成面试练习报告。

        :param interview_results: 面试结果列表 [{question, answer, score_data}]
        :param target_role: 目标岗位
        :param output_path: 输出文件路径（可选）
        :return: 报告文件路径和总评摘要
        """
        if not interview_results:
            return self._error("缺少面试结果数据", detail="interview_results 不能为空")

        self.validate_input({"target_role": target_role}, ["target_role"])

        self.logger.info(
            f"开始生成面试报告 | 岗位: {target_role} | 题目数: {len(interview_results)}"
        )

        if not output_path:
            output_dir = os.path.join(
                os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output"
            )
            os.makedirs(output_dir, exist_ok=True)
            safe_role = re.sub(r'[^\w\u4e00-\u9fff\-]', '_', target_role)[:30]
            output_path = os.path.join(output_dir, f"{safe_role}_面试练习报告.docx")

        summary = self._compute_summary(interview_results)
        content_md = self._build_markdown_report(interview_results, target_role, summary)

        file_path = self._write_report(content_md, output_path, target_role)

        self.logger.info(f"面试报告已生成: {file_path}")

        return self._success(
            data={
                "file_path": file_path,
                "summary": summary,
                "total_questions": len(interview_results),
            },
            message=f"面试练习报告已生成: {file_path}"
        )

    def _compute_summary(self, results: List[Dict]) -> Dict:
        """计算整体汇总"""
        total_score = 0
        dimension_totals: Dict[str, int] = {}
        count = len(results)

        for r in results:
            score_data = r.get("score_data", {})
            total_score += score_data.get("score", 0)

            for dim_name, dim_data in score_data.get("dimensions", {}).items():
                if isinstance(dim_data, dict):
                    dim_score = dim_data.get("score", 0)
                else:
                    dim_score = int(dim_data)
                dimension_totals[dim_name] = dimension_totals.get(dim_name, 0) + dim_score

        avg_score = round(total_score / count) if count else 0
        avg_dimensions = {
            k: round(v / count) for k, v in dimension_totals.items()
        } if count else {}

        tier_map = {90: "S", 75: "A", 60: "B", 40: "C"}
        tier = "D"
        for threshold, t in tier_map.items():
            if avg_score >= threshold:
                tier = t
                break

        all_feedback = []
        for r in results:
            score_data = r.get("score_data", {})
            all_feedback.extend(score_data.get("feedback", []))

        unique_feedback = list(dict.fromkeys(all_feedback))[:5]

        return {
            "average_score": avg_score,
            "tier": tier,
            "average_dimensions": avg_dimensions,
            "top_suggestions": unique_feedback,
            "strongest": max(avg_dimensions, key=avg_dimensions.get) if avg_dimensions else "N/A",
            "weakest": min(avg_dimensions, key=avg_dimensions.get) if avg_dimensions else "N/A",
        }

    def _build_markdown_report(self, results: List[Dict], role: str,
                               summary: Dict) -> str:
        """构建 Markdown 格式报告"""
        lines: List[str] = [
            f"# 面试练习报告",
            f"",
            f"**目标岗位**: {role}",
            f"**练习题数**: {len(results)}",
            f"**综合得分**: {summary['average_score']}/100（等级 {summary['tier']}）",
            f"",
            f"---",
            f"",
            f"## 维度分析",
            f"",
            f"| 维度 | 平均分 |",
            f"|------|--------|",
        ]

        dim_names = {
            "completeness": "内容完整度",
            "logic": "逻辑条理性",
            "star_structure": "STAR 结构性",
            "keyword_coverage": "关键词覆盖",
        }
        for dim, score in summary.get("average_dimensions", {}).items():
            lines.append(f"| {dim_names.get(dim, dim)} | {score}/100 |")

        lines.extend([
            f"",
            f"**最强维度**: {dim_names.get(summary['strongest'], summary['strongest'])}",
            f"**最弱维度**: {dim_names.get(summary['weakest'], summary['weakest'])}",
            f"",
            f"---",
            f"",
            f"## 改进建议",
            f"",
        ])
        for i, fb in enumerate(summary.get("top_suggestions", []), 1):
            lines.append(f"{i}. {fb}")

        lines.extend([
            f"",
            f"---",
            f"",
            f"## 逐题详情",
            f"",
        ])

        for r in results:
            q = r.get("question", "")
            a = r.get("answer", "")
            score_data = r.get("score_data", {})
            score = score_data.get("score", 0)

            lines.extend([
                f"### Q{r.get('id', '?')}: {q}",
                f"",
                f"**得分**: {score}/100",
                f"",
                f"**我的回答**:",
                f"> {a[:500]}",
                f"",
            ])

            feedback = score_data.get("feedback", [])
            if feedback:
                lines.append("**改进建议**:")
                for fb in feedback:
                    lines.append(f"- {fb}")
                lines.append("")

            improved = score_data.get("improved_answer", "")
            if improved:
                lines.extend([
                    f"<details>",
                    f"<summary>查看示范回答</summary>",
                    f"",
                    improved,
                    f"",
                    f"</details>",
                    f"",
                ])

        return "\n".join(lines)

    def _write_report(self, content_md: str, output_path: str,
                      target_role: str) -> str:
        """将报告写入文件（Word 或 Markdown 降级）"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            style = doc.styles["Normal"]
            style.font.name = "微软雅黑"
            style.font.size = Pt(11)

            title = doc.add_heading(f"{target_role} 面试练习报告", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for line in content_md.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                elif line.startswith("# "):
                    pass
                elif line.startswith("## "):
                    doc.add_heading(line.replace("## ", ""), level=2)
                elif line.startswith("### "):
                    doc.add_heading(line.replace("### ", ""), level=3)
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip("*"))
                    run.bold = True
                elif line.startswith("> "):
                    p = doc.add_paragraph(line.replace("> ", ""))
                    p.style = doc.styles["Normal"]
                    p.paragraph_format.left_indent = Inches(0.3)
                elif line.startswith("| ") and "---" not in line:
                    cells = [c.strip() for c in line.split("|")[1:-1]]
                    doc.add_paragraph("  |  ".join(cells))
                elif line.startswith("- "):
                    doc.add_paragraph(line.replace("- ", "  • "), style="List Bullet")
                elif line.startswith(tuple("123456789")) and ". " in line[:4]:
                    doc.add_paragraph(line, style="List Number")
                else:
                    doc.add_paragraph(line)

            doc.save(output_path)
            return output_path

        except ImportError:
            md_path = output_path.replace(".docx", ".md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(content_md)
            self.logger.info("python-docx 不可用，已降级为 Markdown 格式")
            return md_path
