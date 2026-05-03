"""
排版适配 Skill (Resume Writer Skill)
AutoClaw Skill ID: resume_writer

职责：
  接收经过 STAR 打磨的 Markdown 内容和用户 Profile 数据，
  将其填充进预设的标准化简历模板中，
  生成排版精美的 Word (.docx) 简历文件并导出到本地 output/ 目录。

跨应用能力:
  GLM (内容) → Microsoft Word (.docx 文件)

触发场景（通常由 Pipeline 自动触发）：
  - 在 star_polisher 执行完毕后自动触发
  - 或当用户说"帮我生成简历文件 / 导出 Word"

GLM 调度方式：
  GLM 将 polished_md、user_profile、target_role 传入此 Skill。
"""

import os
import re
import json
from typing import Dict, Any
from skills import AutoClawSkill

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResumeWriterSkill(AutoClawSkill):
    """
    排版适配 Skill
    将打磨后的 Markdown 内容填充进 Word 模板，输出 .docx 简历文件
    """

    SKILL_NAME = "resume_writer"
    SKILL_DESCRIPTION = "将 STAR 打磨后的内容填充进简历模板，导出排版精美的 Word 文件至 output/ 目录。"

    def run(self, polished_md: str, target_role: str = "目标岗位",
            user_profile: Dict = None, output_dir: str = None, **kwargs) -> Dict[str, Any]:
        """
        执行简历生成。

        :param polished_md: star_polisher 输出的 Markdown 字符串
        :param target_role: 目标岗位名称（用于命名文件）
        :param user_profile: 用户基本信息字典
        :param output_dir: 输出目录路径（默认 output/）
        :return: 生成的文件路径
        """
        self.validate_input({"polished_md": polished_md}, ["polished_md"])

        # 加载用户 Profile
        if user_profile is None:
            user_profile = self._load_default_profile()

        # 确定输出路径
        if output_dir is None:
            output_dir = os.path.join(
                os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                "output"
            )
        os.makedirs(output_dir, exist_ok=True)

        name = user_profile.get("name", "用户")
        safe_role = re.sub(r'[^\w\u4e00-\u9fff\-]', '_', target_role)[:30]
        output_path = os.path.join(output_dir, f"【{safe_role}】{name}_定向简历.docx")

        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx 未安装，以 Markdown 格式保存简历")
            md_path = output_path.replace(".docx", ".md")
            md_content = self._build_markdown_resume(polished_md, user_profile, target_role)
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_content)
            return self._success(
                data={"file_path": md_path, "format": "markdown"},
                message=f"简历已以 Markdown 格式生成（请安装 python-docx 以生成 Word 格式）"
            )

        self.logger.info(f"开始生成 Word 简历 | 目标岗位：{target_role} | 姓名：{name}")
        self._build_word_resume(polished_md, user_profile, target_role, output_path)

        return self._success(
            data={"file_path": output_path, "format": "docx", "target_role": target_role},
            message=f"Word 简历已生成: {output_path}"
        )

    def _load_default_profile(self) -> Dict:
        """尝试从 assets/user_resume.json 加载用户数据"""
        assets_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "assets", "user_resume.json"
        )
        try:
            with open(assets_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except FileNotFoundError:
            self.logger.warning("未找到 user_resume.json，使用默认空白 Profile")
            return {
                "name": "用户",
                "phone": "待补充",
                "email": "待补充",
                "education": {
                    "school": "XX 大学",
                    "major": "XX 专业",
                    "degree": "本科",
                    "period": "2020.09 - 2024.06"
                },
                "skills": [],
                "target_role": "实习生"
            }

    def _build_word_resume(self, polished_md: str, user: Dict, target_role: str, output_path: str):
        """构建并保存 Word 文档"""
        doc = Document()

        # ── 页面边距 ──
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # ── 姓名标题 ──
        name_para = doc.add_paragraph()
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_run = name_para.add_run(user.get("name", "姓名"))
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # ── 联系方式 ──
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_text = (
            f"Tel: {user.get('phone', '138-XXXX-XXXX')}  "
            f"| Email: {user.get('email', 'your@email.com')}  "
            f"| 求职意向：{target_role}  "
            f"| 期望城市：{user.get('target_city', '上海')}"
        )
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        self._add_divider(doc)

        # ── 教育背景 ──
        self._add_section_heading(doc, "教育背景")
        edu = user.get("education", {})
        edu_para = doc.add_paragraph()
        edu_run = edu_para.add_run(
            f"{edu.get('school', 'XX 大学')}  |  "
            f"{edu.get('major', 'XX 专业')}  |  "
            f"{edu.get('degree', '本科')}  |  "
            f"{edu.get('period', '2020.09 - 2024.06')}"
        )
        edu_run.font.size = Pt(11)

        self._add_divider(doc)

        # ── 核心经历（STAR 打磨内容）──
        self._add_section_heading(doc, "核心经历")
        self._render_markdown_block(doc, polished_md)

        self._add_divider(doc)

        # ── 技能与证书 ──
        self._add_section_heading(doc, "技能与证书")
        skills = user.get("skills", [])
        if skills:
            skills_para = doc.add_paragraph(style="List Bullet")
            skills_para.add_run("专业技能：").bold = True
            skills_para.add_run(", ".join(skills))
        else:
            doc.add_paragraph("专业技能：待补充", style="List Bullet")

        doc.save(output_path)
        self.logger.info(f"Word 简历已保存: {output_path}")

    def _add_section_heading(self, doc, text: str):
        """添加带蓝色背景效果的区块标题"""
        heading = doc.add_heading(text, level=1)
        heading.runs[0].font.size = Pt(13)
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    def _add_divider(self, doc):
        """添加分隔线段落"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "BFBFBF")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_markdown_block(self, doc, md_text: str):
        """将 Markdown 格式的经历文本渲染为 Word 段落"""
        lines = md_text.strip().split("\n")
        for line in lines:
            if not line.strip():
                continue
            if line.startswith("### "):
                # 经历子标题
                exp_heading = doc.add_heading(line.replace("### ", ""), level=2)
                exp_heading.runs[0].font.size = Pt(11.5)
                exp_heading.runs[0].font.color.rgb = RGBColor(0x2E, 0x4F, 0x7B)
            elif line.startswith("- ") or line.startswith("* "):
                # 要点列表
                bullet_text = line[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                # 解析 **加粗** 片段
                self._render_bold_runs(p, bullet_text)
            else:
                doc.add_paragraph(line.strip())

    def _render_bold_runs(self, paragraph, text: str):
        """解析 **text** 加粗标记并应用到 Word run"""
        parts = text.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.font.size = Pt(10.5)
            if i % 2 == 1:  # 奇数索引为加粗内容
                run.bold = True

    def _build_markdown_resume(self, polished_md: str, user: Dict, target_role: str) -> str:
        """降级方案：生成 Markdown 格式简历"""
        edu = user.get("education", {})
        skills = ", ".join(user.get("skills", ["待补充"]))
        return f"""# {user.get('name', '用户')} 的个人简历

📞 {user.get('phone', '待补充')} | ✉️ {user.get('email', '待补充')} | 🎯 求职意向：{target_role} | 📍 {user.get('target_city', '上海')}

---

## 🎓 教育背景
**{edu.get('school', 'XX大学')}** | {edu.get('major', 'XX专业')} | {edu.get('degree', '本科')} | {edu.get('period', '2020.09 - 2024.06')}

---

## 💼 核心经历

{polished_md}

---

## 🛠️ 技能与证书
- 专业技能：{skills}

> ⚠️ 提示：请安装 `python-docx` 以生成正式的 Word 格式简历
"""
