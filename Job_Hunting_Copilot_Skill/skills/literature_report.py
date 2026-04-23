"""
文献综述报告生成 Skill (Literature Report Skill)
AutoClaw Skill ID: literature_report

职责：
  接收 paper_digest 输出的结构化文献数据，
  生成排版精美的 Word (.docx) 文献综述报告，
  包含论文表格、趋势分析、研究空白建议。

跨应用能力:
  GLM (提炼数据) -> Microsoft Word (.docx 文件)

触发场景（通常由 Pipeline 自动触发）：
  - 在 paper_digest 执行完毕后自动触发
"""

import os
from typing import List, Dict, Any, Optional
from skills import AutoClawSkill

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class LiteratureReportSkill(AutoClawSkill):
    """
    文献综述报告生成 Skill
    将结构化文献数据生成 Word 报告
    """

    SKILL_NAME = "literature_report"
    SKILL_DESCRIPTION = "将文献提炼数据生成 Word 文献综述报告，含论文表格、趋势分析、研究建议。"

    def run(self, digested_papers: List[Dict], review_md: str = "",
            trends: List[str] = None, topic: str = "",
            **kwargs) -> Dict[str, Any]:
        """
        生成文献综述 Word 报告。

        :param digested_papers: paper_digest 输出的结构化论文列表
        :param review_md: 文献综述 Markdown（可选）
        :param trends: 趋势分析列表（可选）
        :param topic: 研究主题
        :return: 生成的文件路径
        """
        self.validate_input({"digested_papers": digested_papers}, ["digested_papers"])

        output_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "output"
        )
        os.makedirs(output_dir, exist_ok=True)

        safe_topic = topic.replace(" ", "_").replace("/", "_")[:30] if topic else "文献综述"
        output_path = os.path.join(output_dir, f"{safe_topic}_文献综述报告.docx")

        if not DOCX_AVAILABLE:
            md_path = output_path.replace(".docx", ".md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(review_md or f"# {topic} 文献综述\n\n请安装 python-docx 生成 Word 报告")
            return self._success(
                data={"file_path": md_path, "format": "markdown"},
                message=f"文献综述已以 Markdown 格式保存"
            )

        self._build_word_report(digested_papers, trends or [], topic, output_path)

        return self._success(
            data={"file_path": output_path, "format": "docx", "paper_count": len(digested_papers)},
            message=f"文献综述报告已生成: {output_path}"
        )

    def _build_word_report(self, papers: List[Dict], trends: List[str],
                           topic: str, output_path: str):
        """构建 Word 文档"""
        doc = Document()

        # 页面边距
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.8)

        # 标题
        title = doc.add_heading(f"{topic} -- 文献综述报告", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in title.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # 概览
        doc.add_heading("一、调研概览", level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"本报告共调研 {len(papers)} 篇与「{topic}」相关的学术论文。")
        if trends:
            for trend in trends:
                doc.add_paragraph(trend, style="List Bullet")

        # 论文表格
        doc.add_heading("二、论文详情表", level=1)
        self._add_paper_table(doc, papers)

        # 综述内容
        if papers:
            doc.add_heading("三、分篇综述", level=1)
            for i, paper in enumerate(papers, 1):
                h = doc.add_heading(f"{i}. {paper.get('title', '无标题')}", level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

                meta = doc.add_paragraph()
                meta.add_run(f"作者: {paper.get('authors', '未知')} | "
                             f"年份: {paper.get('year', 'N/A')} | "
                             f"引用: {paper.get('citations', 0)}").font.size = Pt(9)

                methods = paper.get("core_methods", [])
                if methods:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run("方法: ").bold = True
                    p.add_run(", ".join(methods))

                findings = paper.get("key_findings", [])
                if findings:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run("发现: ").bold = True
                    p.add_run("; ".join(findings[:3]))

                conclusion = paper.get("conclusion", "")
                if conclusion:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run("结论: ").bold = True
                    p.add_run(conclusion)

        # 研究空白建议
        doc.add_heading("四、研究空白与建议", level=1)
        suggestions = self._generate_suggestions(papers, topic)
        for s in suggestions:
            doc.add_paragraph(s, style="List Bullet")

        doc.save(output_path)
        self.logger.info(f"文献综述报告已保存: {output_path}")

    def _add_paper_table(self, doc, papers: List[Dict]):
        """添加论文概览表格"""
        headers = ["#", "标题", "作者", "年份", "引用数", "核心方法"]
        table = doc.add_table(rows=1 + len(papers), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        HEADER_BG = "1F4E79"
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), HEADER_BG)
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)

        for r_idx, paper in enumerate(papers):
            methods = ", ".join(paper.get("core_methods", [])[:2])
            row_data = [
                str(r_idx + 1),
                paper.get("title", "无标题")[:40],
                paper.get("authors", "未知")[:20],
                str(paper.get("year", "N/A")),
                str(paper.get("citations", 0)),
                methods[:25],
            ]
            for c_idx, val in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ""
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(8)

        # 设置列宽
        widths = [1, 6, 3.5, 1.5, 1.5, 3.5]
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

        doc.add_paragraph()

    def _generate_suggestions(self, papers: List[Dict], topic: str) -> List[str]:
        """生成研究空白与建议"""
        suggestions = [
            f"当前关于「{topic}」的研究主要集中在方法层面，实际落地应用的长期效果评估仍不足",
            "跨领域迁移能力的研究较为匮乏，建议关注通用性更强的解决方案",
            "数据集规模和多样性有待提升，建议构建更大规模的中文领域基准测试",
        ]
        if len(papers) < 5:
            suggestions.append(f"当前仅检索到 {len(papers)} 篇相关论文，建议扩大检索范围或调整关键词")
        return suggestions
